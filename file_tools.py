"""Utility helpers for text extraction from uploaded files."""

from __future__ import annotations

import io
import logging

from PyPDF2 import PdfReader
from fpdf import FPDF
import docx
from docx.shared import Inches


def extract_text_from_file(file_data: bytes, file_type: str) -> str:
    """Return raw text from a PDF or DOCX file.

    Args:
        file_data: File content as byte string.
        file_type: MIME type of the uploaded file.

    Returns:
        Concatenated text extracted from all pages or paragraphs.
    """
    text = ""
    try:
        if file_type == "application/pdf":
            reader = PdfReader(io.BytesIO(file_data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        elif file_type in (
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            doc = docx.Document(io.BytesIO(file_data))
            text = "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception as exc:  # pragma: no cover - log only
        logging.error("Error reading file: %s", exc)
    return text


def create_pdf(text: str, *, font: str = "Arial", logo: bytes | None = None) -> bytes:
    """Return text as PDF bytes with optional logo and font."""

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    try:
        pdf.set_font(font, size=12)
    except Exception:
        pdf.set_font("Arial", size=12)

    if logo:
        with io.BytesIO(logo) as buf:
            pdf.image(buf, x=10, y=8, w=30)
        pdf.ln(25)

    for line in text.splitlines():
        pdf.multi_cell(0, 10, text=line)
    return bytes(pdf.output())


def create_docx(text: str, *, font: str = "Arial", logo: bytes | None = None) -> bytes:
    """Return text as Word bytes with optional logo and font."""

    doc = docx.Document()
    style = doc.styles["Normal"].font
    style.name = font

    if logo:
        with io.BytesIO(logo) as buf:
            doc.add_picture(buf, width=Inches(1.0))

    for line in text.splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
