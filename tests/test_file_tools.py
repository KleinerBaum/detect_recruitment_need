import io
from pathlib import Path

import importlib.util
import docx
from fpdf import FPDF


def load_file_tools():
    path = Path(__file__).resolve().parents[1] / "file_tools.py"
    spec = importlib.util.spec_from_file_location("file_tools", path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader
    spec.loader.exec_module(module)
    return module


extract_text_from_file = load_file_tools().extract_text_from_file
create_pdf = load_file_tools().create_pdf
create_docx = load_file_tools().create_docx


def test_extract_text_from_pdf(tmp_path: Path) -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, "Hello PDF", new_y="NEXT")
    pdf_bytes = bytes(pdf.output(dest="S"))

    out = extract_text_from_file(pdf_bytes, "application/pdf")
    assert "Hello PDF" in out


def test_extract_text_from_docx(tmp_path: Path) -> None:
    doc = docx.Document()
    doc.add_paragraph("Hello DOCX")
    buf = io.BytesIO()
    doc.save(buf)
    out = extract_text_from_file(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "Hello DOCX" in out


def test_extract_text_handles_errors() -> None:
    assert extract_text_from_file(b"invalid", "application/pdf") == ""
    assert (
        extract_text_from_file(
            b"invalid",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        == ""
    )


def test_create_pdf_and_docx() -> None:
    logo = b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01\x0d\n\x2d\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
    pdf_bytes = create_pdf("Hello PDF", font="Arial", logo=logo)
    assert pdf_bytes.startswith(b"%PDF")

    docx_bytes = create_docx("Hello DOCX", font="Arial", logo=logo)
    doc = docx.Document(io.BytesIO(docx_bytes))
    assert doc.paragraphs[-1].text == "Hello DOCX"
